"""

Markdown to DOCX conversion module extracted from md2docx-python project.

Project address: https://github.com/shloktech/md2docx-python/
Project LICENSE: LICENSE.MD2DOCX

The reason for extraction is to align the API and features with the needs.
Changes:
- input is text, not file
- markdown2 is used instead of markdown
- table support is added via markdown2 extras and additional HTML processing

"""

import markdown2
import logging
import os
from docx import Document
from docx.shared import Inches
from bs4 import BeautifulSoup, Tag

IMAGE_WIDTH_IN_INCHES = 6


def get_element_text(element: Tag) -> str:
    if hasattr(element, "get_text"):
        return element.get_text(strip=True)
    else:
        return str(element).strip()


def process_list_items(list_element: Tag, doc: Document, style_base: str, level=0):
    # Get direct children li elements only (not nested)
    for li in list_element.find_all("li", recursive=False):
        # Get text content, excluding nested lists
        text_parts = []
        for child in li.children:
            if child.name not in ["ul", "ol"]:
                text_parts.append(get_element_text(child))

        text = " ".join(text_parts).strip()

        # Add paragraph with appropriate indentation level
        if text:
            style = style_base if level == 0 else f"{style_base} {level + 1}"
            doc.add_paragraph(text, style=style)

        # Process nested lists
        nested_ul = li.find("ul", recursive=False)
        nested_ol = li.find("ol", recursive=False)

        if nested_ul:
            process_list_items(nested_ul, doc, "List Bullet", level + 1)
        if nested_ol:
            process_list_items(nested_ol, doc, "List Number", level + 1)


def embed_image(img: Tag, doc: Document, base_path: str = ""):
    img_src = img.get("src")
    img_title = img.get("title", "").strip()
    img_alt = img.get("alt", "").strip()

    # Use title if available, otherwise use alt text
    caption_text = img_title if img_title else img_alt

    if img_src:
        # Try the image path as-is first, then relative to base_path
        image_path = img_src
        if not os.path.exists(image_path) and base_path:
            image_path = os.path.join(base_path, img_src)

        if os.path.exists(image_path):
            try:
                doc.add_picture(image_path, width=Inches(IMAGE_WIDTH_IN_INCHES))
                if caption_text:
                    caption_paragraph = doc.add_paragraph(caption_text)
                    caption_paragraph.style = "Caption"
            except Exception as e:
                logging.error(f"Exception while adding image {e}")
                pass


def markdown_to_word_file(
    markdown_source: str, word_file_path: str, base_path: str = ""
):
    doc = markdown_to_word_object(markdown_source, base_path)
    doc.save(word_file_path)


def markdown_to_word_object(markdown_source: str, base_path: str = "") -> Document:
    # Converting Markdown to HTML
    html_content = markdown2.markdown(markdown_source, extras=["tables", "wiki-tables"])

    # Creating a new Word Document
    doc = Document()

    # Converting HTML to text and adding it to the Word Document
    soup = BeautifulSoup(html_content, "html.parser")

    # Adding content to the Word Document
    for element in soup:
        if element.name == "h1":
            doc.add_heading(element.text, level=1)
        elif element.name == "h2":
            doc.add_heading(element.text, level=2)
        elif element.name == "h3":
            doc.add_heading(element.text, level=3)
        elif element.name == "p":
            # Check if paragraph contains an image
            img = element.find("img")
            if img:
                embed_image(img, doc, base_path)
            else:
                # Regular paragraph without image
                paragraph = doc.add_paragraph()
                for child in element.children:
                    if child.name == "strong":
                        paragraph.add_run(child.text).bold = True
                    elif child.name == "em":
                        paragraph.add_run(child.text).italic = True
                    else:
                        paragraph.add_run(child)
        elif element.name == "ul":
            process_list_items(element, doc, "List Bullet")
        elif element.name == "ol":
            process_list_items(element, doc, "List Number")
        elif element.name == "table":
            rows_data = []
            for row in element.find_all("tr"):
                cells = row.find_all(["th", "td"])
                row_data = [cell.get_text(strip=True) for cell in cells]
                if row_data:
                    rows_data.append(row_data)

            if rows_data:
                columns_count = len(rows_data[0])
                table = doc.add_table(rows=len(rows_data), cols=columns_count)
                table.style = "Table Grid"

                for row_index, row_data in enumerate(rows_data):
                    for column_index, cell_text in enumerate(row_data):
                        if column_index < columns_count:
                            table.rows[row_index].cells[column_index].text = cell_text

                # Make the first row bold if it is a header
                first_row = element.find("tr")
                if first_row and first_row.find("th"):
                    for cell in table.rows[0].cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

    return doc
