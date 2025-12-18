"""
Tests for md2docx module
"""

import pytest
from docx.document import Document as DocumentType
from templateprocessor.md2docx import markdown_to_word_object


class TestMarkdownToWordObject:
    """
    Test cases for markdown_to_word_object function.
    markdown_to_word_file is not tested, as it is a simple file write.
    """

    def test_simple_text(self):
        """Test converting simple text paragraphs."""
        # Prepare
        markdown = "This is a simple paragraph.\n\nThis is another paragraph."

        # Execute
        doc = markdown_to_word_object(markdown)
        # Verify
        assert isinstance(doc, DocumentType)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        assert len(paragraphs) == 2
        assert "This is a simple paragraph." in paragraphs[0]
        assert "This is another paragraph." in paragraphs[1]

    def test_simple_list(self):
        """Test converting a simple bullet list."""
        # Prepare
        markdown = """
- First item
- Second item
- Third item
"""

        # Execute
        doc = markdown_to_word_object(markdown)
        # Verify
        assert isinstance(doc, DocumentType)
        paragraphs = [p for p in doc.paragraphs if p.text.strip()]
        assert len(paragraphs) == 3
        assert paragraphs[0].text == "First item"
        assert paragraphs[1].text == "Second item"
        assert paragraphs[2].text == "Third item"
        # Verify style
        assert "List Bullet" in paragraphs[0].style.name
        assert "List Bullet" in paragraphs[1].style.name
        assert "List Bullet" in paragraphs[2].style.name

    def test_nested_list_two_levels(self):
        """Test converting a nested list with 2 levels."""
        # Prepare
        markdown = """
- Top level item 1
  - Nested item 1.1
  - Nested item 1.2
- Top level item 2
  - Nested item 2.1
"""

        # Execute
        doc = markdown_to_word_object(markdown)
        # Verify
        assert isinstance(doc, DocumentType)
        paragraphs = [p for p in doc.paragraphs if p.text.strip()]
        assert len(paragraphs) == 5

        # Check text content
        assert paragraphs[0].text == "Top level item 1"
        assert paragraphs[1].text == "Nested item 1.1"
        assert paragraphs[2].text == "Nested item 1.2"
        assert paragraphs[3].text == "Top level item 2"
        assert paragraphs[4].text == "Nested item 2.1"

        # Verify top level uses base style
        assert "List Bullet" in paragraphs[0].style.name
        assert "List Bullet" in paragraphs[3].style.name

        # Verify nested items use appropriate style
        assert "List Bullet 2" in paragraphs[1].style.name
        assert "List Bullet 2" in paragraphs[2].style.name

    def test_table(self):
        """Test converting a markdown table."""
        # Prepare
        markdown = """
| Header 1 | Header 2 | Header 3 |
|----------|----------|----------|
| Row 1 Col 1 | Row 1 Col 2 | Row 1 Col 3 |
| Row 2 Col 1 | Row 2 Col 2 | Row 2 Col 3 |
"""

        # Execute
        doc = markdown_to_word_object(markdown)
        # Verify
        assert isinstance(doc, DocumentType)
        assert len(doc.tables) == 1

        table = doc.tables[0]
        assert len(table.rows) == 3
        assert len(table.columns) == 3

        # Check header row
        assert table.rows[0].cells[0].text == "Header 1"
        assert table.rows[0].cells[1].text == "Header 2"
        assert table.rows[0].cells[2].text == "Header 3"

        # Check data rows
        assert table.rows[1].cells[0].text == "Row 1 Col 1"
        assert table.rows[1].cells[1].text == "Row 1 Col 2"
        assert table.rows[1].cells[2].text == "Row 1 Col 3"

        assert table.rows[2].cells[0].text == "Row 2 Col 1"
        assert table.rows[2].cells[1].text == "Row 2 Col 2"
        assert table.rows[2].cells[2].text == "Row 2 Col 3"

        # Verify header row is bold
        first_cell_runs = table.rows[0].cells[0].paragraphs[0].runs
        assert first_cell_runs is not None
        assert first_cell_runs[0].bold

    def test_header(self):
        """Test converting markdown headers."""
        # Prepare
        markdown = """
# Header 1

## Header 2

### Header 3
"""

        # Execute
        doc = markdown_to_word_object(markdown)

        # Verify
        assert isinstance(doc, DocumentType)
        paragraphs = [p for p in doc.paragraphs if p.text.strip()]
        assert len(paragraphs) == 3

        # Check header text
        assert paragraphs[0].text == "Header 1"
        assert paragraphs[1].text == "Header 2"
        assert paragraphs[2].text == "Header 3"

        # Verify heading styles
        assert "Heading 1" in paragraphs[0].style.name
        assert "Heading 2" in paragraphs[1].style.name
        assert "Heading 3" in paragraphs[2].style.name
